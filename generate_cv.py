#!/usr/bin/env python3
"""Generate Santosh Jammi CV in Word (.docx) format."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── PAGE MARGINS ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── COLOUR PALETTE ────────────────────────────────────────────────
INDIGO   = RGBColor(0x6B, 0x68, 0xF0)   # #6B68F0 — accent
DARK     = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A — headings
MID      = RGBColor(0x47, 0x55, 0x69)   # #475569 — sub-text
LIGHT    = RGBColor(0x94, 0xA3, 0xB8)   # #94A3B8 — muted
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
RULE_CLR = RGBColor(0xE2, 0xE8, 0xF0)   # #E2E8F0 — divider

# ── HELPER: add a horizontal rule ─────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── HELPER: section heading ────────────────────────────────────────
def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = INDIGO
    run.font.name = 'Calibri'
    add_hr(doc)

# ── HELPER: bullet point ──────────────────────────────────────────
def add_bullet(doc, label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25)
    if label:
        r = p.add_run(f'[{label}] ')
        r.bold = True
        r.font.color.rgb = INDIGO
        r.font.size = Pt(9.5)
        r.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK
    r2.font.name = 'Calibri'

# ── HELPER: plain paragraph ───────────────────────────────────────
def add_para(doc, text, size=9.5, color=DARK, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.size   = Pt(size)
    r.font.color.rgb = color
    r.font.name   = 'Calibri'
    return p

# ═══════════════════════════════════════════════════════════════════
#  HEADER  — Name + Title
# ═══════════════════════════════════════════════════════════════════
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_before = Pt(0)
name_p.paragraph_format.space_after  = Pt(4)
name_r = name_p.add_run('Santosh Jammi')
name_r.bold = True
name_r.font.size = Pt(22)
name_r.font.color.rgb = DARK
name_r.font.name = 'Calibri'

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
title_r = title_p.add_run(
    'Senior Data Engineer → Principal Data Architect\n'
    'Microsoft Fabric & Azure Native Ecosystems | Cloud FinOps Specialist'
)
title_r.font.size = Pt(10.5)
title_r.font.color.rgb = INDIGO
title_r.font.name = 'Calibri'

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(4)
contact_r = contact_p.add_run(
    'Hyderabad, India  ·  santu.jammi@gmail.com  ·  +91-9110580037  ·  linkedin.com/in/santoshjammi29  ·  github.com/santoshjammi29'
)
contact_r.font.size = Pt(9)
contact_r.font.color.rgb = MID
contact_r.font.name = 'Calibri'

# Certifications strip
cert_p = doc.add_paragraph()
cert_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cert_p.paragraph_format.space_after = Pt(6)
certs = [
    'DP-600 Fabric Analytics Engineer', 'DP-203 Azure Data Engineer',
    'PL-300 Power BI Analyst', 'AI-900 Azure AI Fundamentals',
    'DP-900 Data Fundamentals', 'AZ-900 Azure Fundamentals',
]
cert_r = cert_p.add_run('  |  '.join(certs))
cert_r.font.size = Pt(8.5)
cert_r.font.color.rgb = INDIGO
cert_r.bold = True
cert_r.font.name = 'Calibri'

add_hr(doc)

# ═══════════════════════════════════════════════════════════════════
#  PROFESSIONAL SUMMARY
# ═══════════════════════════════════════════════════════════════════
add_section_heading(doc, 'Professional Summary')
summary_p = doc.add_paragraph()
summary_p.paragraph_format.space_after = Pt(4)
summary_r = summary_p.add_run(
    'Strategic Lead Data Engineer with 8+ years of experience and a proven 5-year track record of leading '
    'high-performance engineering teams and managing direct stakeholder engagements. Expert in designing '
    'scalable Lakehouse architectures on Microsoft Azure and Microsoft Fabric, acting as the primary '
    'techno-functional bridge between business goals and technical delivery. Specialises in Cloud FinOps, '
    'significantly reducing operational costs through advanced compute optimisation (DIUs, Parallelism) and '
    'SLA optimisation. Deep domain expertise in the EU BFSI sector, ensuring strict GDPR / data sovereignty '
    'compliance while delivering rapid innovation and 99.9% system reliability.'
)
summary_r.font.size = Pt(9.5)
summary_r.font.color.rgb = DARK
summary_r.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════════
#  CORE TECHNICAL COMPETENCIES
# ═══════════════════════════════════════════════════════════════════
add_section_heading(doc, 'Core Technical Competencies')

skills_data = [
    ('Core Architecture',
     'Medallion Architecture, Data Mesh, Cloud FinOps, Disaster Recovery (DR), Governance, '
     'Stakeholder Management, Kimball Modeling, Inmon Methodology, Data Vault 2.0, '
     'Star / Snowflake Schemas, SCDs, Zero-Copy Architecture'),
    ('Microsoft Fabric & Azure',
     'Microsoft Fabric, OneLake & Direct Lake, Azure Synapse Analytics, Azure Data Factory (ADF), '
     'ADLS Gen2, Stream Analytics, Event Hubs, Databricks, Azure SQL, Azure Cosmos DB, '
     'Logic Apps, Azure Monitor / KQL'),
    ('ETL / ELT & Pipelines',
     'ADF Mapping Data Flows, Idempotent Pipeline Design, Apache Airflow, SCD Type 2, '
     'ACID Compliance, Delta Lake, Liquid Clustering, Z-Ordering, Copy-On-Write (CoW), '
     'MERGE INTO / UPSERT, Delta Time Travel, Data Orchestration'),
    ('DevOps & Governance',
     'Azure DevOps, CI/CD (YAML), Git Branching, Automated Rollbacks, Row-Level Security (RLS), '
     'GDPR Compliance, Data Sovereignty, Azure Key Vault, Service Principals, OAuth Handshake, '
     'Entra ID, RBAC, T-SQL Advanced, KQL, ARM Templates, Schema Evolution, Data Quality Management'),
]

for category, items in skills_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f'{category}: ')
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = DARK
    r1.font.name = 'Calibri'
    r2 = p.add_run(items)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = MID
    r2.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════════
#  PROFESSIONAL EXPERIENCE
# ═══════════════════════════════════════════════════════════════════
add_section_heading(doc, 'Professional Experience')

# ── TCS ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(1)
r = p.add_run('Senior Data Engineer')
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(1)
r2a = p2.add_run('Tata Consultancy Services — EU BFSI / Enterprise')
r2a.font.size = Pt(9.5); r2a.font.color.rgb = INDIGO; r2a.bold = True; r2a.font.name = 'Calibri'
r2b = p2.add_run('     Hyderabad, India     |     Apr 2022 – Present')
r2b.font.size = Pt(9); r2b.font.color.rgb = MID; r2b.font.name = 'Calibri'

overview_p = doc.add_paragraph()
overview_p.paragraph_format.space_after = Pt(2)
r_ov = overview_p.add_run(
    'Strategic engineering leader serving as the primary onsite technical point of contact for Product Owners '
    'and Architects. Managed offshore delivery of 5+ engineers, ensuring zero schedule slippage across sprint '
    'deliverables while building highly resilient systems.'
)
r_ov.italic = True; r_ov.font.size = Pt(9.5); r_ov.font.color.rgb = MID; r_ov.font.name = 'Calibri'

tcs_bullets = [
    ('Leadership',  'Served as primary technical point of contact for onsite Product Owners and Architects, facilitating requirement discovery, managing offshore delivery of 5+ engineers, and ensuring zero schedule slippage across all sprint deliverables.'),
    ('Lakehouse',   'Designed a robust Medallion Architecture (Bronze / Silver / Gold) managing 5TB+ data. Implemented idempotent ETL patterns with strict integrity constraints (MERGE INTO / ACID Delta Lake) to ensure 100% data consistency across all pipeline retries and transient cluster failures.'),
    ('FinOps',      'Led compute optimisation — dynamically tuning DIUs and Parallelism in Azure Data Factory and resolving Integration Runtime bottlenecks — delivering a 40% reduction in data processing latency and significant cloud cost savings across enterprise pipelines.'),
    ('Fabric',      'Architected resilient Microsoft Fabric semantic models using Direct Lake mode, carefully managing F64 capacity memory guardrails (25 GB) to entirely eliminate refresh failures and ensure high availability for C-level dashboards, preventing DirectQuery fallback latency.'),
    ('DevOps',      'Automated CI/CD via Azure DevOps / YAML with automated rollback strategies; enforced strict Schema Evolution policies to maintain 100% downstream platform stability across 5+ concurrent engineering squads.'),
]
for label, text in tcs_bullets:
    add_bullet(doc, label, text)

# ── LTIMindtree ──────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(1)
r = p.add_run('Senior Data Engineer')
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(1)
r2a = p2.add_run('Larsen & Toubro Infotech (LTIMindtree) — Enterprise Data')
r2a.font.size = Pt(9.5); r2a.font.color.rgb = INDIGO; r2a.bold = True; r2a.font.name = 'Calibri'
r2b = p2.add_run('     Bengaluru, India     |     Apr 2021 – Apr 2022')
r2b.font.size = Pt(9); r2b.font.color.rgb = MID; r2b.font.name = 'Calibri'

overview_p = doc.add_paragraph()
overview_p.paragraph_format.space_after = Pt(2)
r_ov = overview_p.add_run(
    'Orchestrated enterprise data modernisation pipelines, translating complex client business logic into '
    'high-efficiency Azure infrastructure.'
)
r_ov.italic = True; r_ov.font.size = Pt(9.5); r_ov.font.color.rgb = MID; r_ov.font.name = 'Calibri'

lti_bullets = [
    ('ETL Performance', 'Re-engineered legacy pipelines using Azure Data Factory, solving the Small File Problem and cutting query execution times by 40% through parallelism tuning, partition optimisation, and Integration Runtime right-sizing.'),
    ('Reporting',       'Redesigned Power BI models with aggregation tables, improving report load speeds by 25% and achieving high availability across all executive and operational reporting workloads.'),
]
for label, text in lti_bullets:
    add_bullet(doc, label, text)

# ── Infosys ──────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(1)
r = p.add_run('Technology Analyst  →  Senior System Engineer  →  System Engineer')
r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(1)
r2a = p2.add_run('Infosys Limited — Enterprise Technology')
r2a.font.size = Pt(9.5); r2a.font.color.rgb = INDIGO; r2a.bold = True; r2a.font.name = 'Calibri'
r2b = p2.add_run('     Hyderabad, India     |     Dec 2017 – Apr 2021')
r2b.font.size = Pt(9); r2b.font.color.rgb = MID; r2b.font.name = 'Calibri'

overview_p = doc.add_paragraph()
overview_p.paragraph_format.space_after = Pt(2)
r_ov = overview_p.add_run(
    'Led end-to-end data platform engineering and BI solution delivery, specialising in secure architecture '
    'design and transactional query tuning for financial reporting workloads.'
)
r_ov.italic = True; r_ov.font.size = Pt(9.5); r_ov.font.color.rgb = MID; r_ov.font.name = 'Calibri'

infosys_bullets = [
    ('Data Quality', 'Improved data model accuracy by 20% through automated data cleansing and validation frameworks, establishing consistent data quality standards across enterprise financial reporting pipelines.'),
    ('Security',     'Integrated Power BI with SaaS applications, implementing Row-Level Security (RLS) for sensitive financial data and designing role-based access controls governing the end-to-end lifecycle from data modelling to dashboard deployment.'),
    ('Performance',  'Increased data processing efficiency by 15% through advanced T-SQL query tuning, index optimisation, and stored procedure refactoring across large-scale SQL Server and Azure SQL workloads.'),
    ('Platform',     'Collaborated with cross-functional teams to develop high-impact data products utilising SQL Server and Azure services for efficient large-scale processing, supporting business intelligence initiatives across multiple domains.'),
]
for label, text in infosys_bullets:
    add_bullet(doc, label, text)

# ═══════════════════════════════════════════════════════════════════
#  KEY PROJECTS & ARCHITECTURAL HIGHLIGHTS
# ═══════════════════════════════════════════════════════════════════
add_section_heading(doc, 'Key Projects & Architectural Highlights')

projects = [
    {
        'title': '5TB Medallion Lakehouse — EU Enterprise (TCS)',
        'tech':  'Fabric · Delta Lake · ADLS Gen2 · ADF · Azure Monitor',
        'desc':  'Designed Bronze-Silver-Gold Medallion architecture managing 5TB+ data for EU BFSI clients at TCS. Implemented idempotent MERGE INTO patterns ensuring 100% data consistency across all pipeline retries. Delivered 99.9% platform reliability with full GDPR compliance.',
    },
    {
        'title': 'Microsoft Fabric Semantic Model — C-Level Dashboards',
        'tech':  'Microsoft Fabric · Direct Lake · OneLake · Power BI · DP-600',
        'desc':  'Architected resilient Fabric semantic models using Direct Lake mode at TCS. Managed F64 capacity guardrails (25 GB memory limit) to eliminate all refresh failures, accelerating C-level dashboard render speeds and achieving high availability for executive reporting.',
    },
    {
        'title': 'ADF Pipeline Re-engineering — Small File Elimination (LTI)',
        'tech':  'Azure Data Factory · Synapse · Power BI · Aggregation Tables',
        'desc':  'Re-engineered legacy ETL pipelines at LTIMindtree to solve the Small File Problem, cutting query execution times by 40%. Simultaneously redesigned Power BI aggregation models, improving report load speeds by 25% across all high-availability business reports.',
    },
    {
        'title': 'FinOps & Compute Optimisation — DIU Tuning (TCS)',
        'tech':  'ADF · DIU Optimisation · Integration Runtime · CI/CD YAML',
        'desc':  'Led Cloud FinOps initiative at TCS optimising DIU allocation and parallelism across enterprise ADF pipelines. Resolved Integration Runtime bottlenecks and automated CI/CD via Azure DevOps YAML, delivering a 40% reduction in data processing latency with significant measurable cost savings.',
    },
]

for proj in projects:
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after  = Pt(1)
    r_t = p_title.add_run(proj['title'])
    r_t.bold = True; r_t.font.size = Pt(9.5); r_t.font.color.rgb = DARK; r_t.font.name = 'Calibri'

    p_tech = doc.add_paragraph()
    p_tech.paragraph_format.space_after = Pt(1)
    r_tech = p_tech.add_run(proj['tech'])
    r_tech.font.size = Pt(8.5); r_tech.font.color.rgb = INDIGO; r_tech.font.name = 'Calibri'

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(4)
    r_desc = p_desc.add_run(proj['desc'])
    r_desc.font.size = Pt(9.5); r_desc.font.color.rgb = MID; r_desc.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════════
#  EDUCATION
# ═══════════════════════════════════════════════════════════════════
add_section_heading(doc, 'Education')

education = [
    ('B.Tech – Electrical & Electronics Engineering (EEE) — 70%',
     'Raghu Engineering College, Visakhapatnam  |  Aug 2013 – May 2017'),
    ('Intermediate – MPC — 90%',
     'Narayana Junior College, Vizianagaram  |  May 2011 – May 2013'),
    ('Class 10th — 90%',
     'New Central School, Vizianagaram  |  Mar 2010 – Mar 2011'),
]

for degree, school in education:
    p_deg = doc.add_paragraph()
    p_deg.paragraph_format.space_before = Pt(4)
    p_deg.paragraph_format.space_after  = Pt(1)
    r_deg = p_deg.add_run(degree)
    r_deg.bold = True; r_deg.font.size = Pt(9.5); r_deg.font.color.rgb = DARK; r_deg.font.name = 'Calibri'
    p_sch = doc.add_paragraph()
    p_sch.paragraph_format.space_after = Pt(2)
    r_sch = p_sch.add_run(school)
    r_sch.font.size = Pt(9); r_sch.font.color.rgb = MID; r_sch.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════
output_path = 'Santosh_Jammi_CV.docx'
doc.save(output_path)
print(f'✅ CV saved to {output_path}')
